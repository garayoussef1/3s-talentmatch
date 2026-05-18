# Genere 5 CV DOCX pour le scenario de test DevOps/Cloud.
# Resultat dans : data/test_scenario/
# Utilisation : .\.venv-10\Scripts\python.exe scripts\generate_test_cvs.py
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "data", "test_scenario")
os.makedirs(OUT_DIR, exist_ok=True)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def section(doc, title, body_lines):
    heading(doc, title, level=2)
    for line in body_lines:
        doc.add_paragraph(line)


def save(doc, filename):
    path = os.path.join(OUT_DIR, filename)
    doc.save(path)
    print(f"  Généré : {path}")


# ──────────────────────────────────────────────────────────────
# CV 1 — Karim Bensalah — EXCELLENT MATCH
# ──────────────────────────────────────────────────────────────
def cv1():
    doc = Document()
    heading(doc, "Karim BENSALAH")
    doc.add_paragraph("Ingénieur DevOps / Cloud Senior | Paris | karim.bensalah@email.fr | +33 6 12 34 56 78")

    section(doc, "Résumé", [
        "Ingénieur DevOps avec 4 ans d'expérience en environnements cloud AWS et infrastructures conteneurisées. "
        "Maîtrise complète de la chaîne CI/CD, de l'orchestration Kubernetes, du provisionnement Terraform "
        "et de l'automatisation Ansible. Passionné par la fiabilité des systèmes et la culture DevOps."
    ])

    section(doc, "Compétences techniques", [
        "Conteneurisation & Orchestration : Docker, Kubernetes (K8s), Helm, Docker Compose",
        "Cloud : AWS (EC2, S3, EKS, RDS, IAM, CloudWatch), notions GCP",
        "Infrastructure as Code : Terraform, Ansible, CloudFormation",
        "CI/CD : GitLab CI, Jenkins, GitHub Actions, ArgoCD",
        "Monitoring : Prometheus, Grafana, ELK Stack (Elasticsearch, Logstash, Kibana)",
        "Scripting : Python, Bash, YAML",
        "OS : Linux (Ubuntu, CentOS, Debian), administration système avancée",
        "Sécurité : gestion des secrets (Vault), RBAC Kubernetes, IAM AWS",
    ])

    section(doc, "Expérience professionnelle", [
        "Ingénieur DevOps Senior — TechCorp Paris (Jan 2022 – Présent)",
        "- Migration complète de l'infrastructure vers AWS EKS avec Terraform",
        "- Mise en place de pipelines CI/CD GitLab CI pour 12 microservices",
        "- Réduction du temps de déploiement de 45 min à 8 min (-82%)",
        "- Automatisation du provisionnement avec Ansible (200+ serveurs)",
        "",
        "Ingénieur DevOps — StartupCloud (Juin 2020 – Déc 2021)",
        "- Containerisation de 8 applications legacy avec Docker",
        "- Déploiement Kubernetes sur cluster auto-géré (kubeadm)",
        "- Monitoring Prometheus/Grafana et alerting PagerDuty",
    ])

    section(doc, "Formation", [
        "Diplôme d'Ingénieur Informatique — EPITA, Paris (2020)",
        "Certifications : AWS Solutions Architect Associate, Certified Kubernetes Administrator (CKA)",
    ])

    section(doc, "Langues", ["Français (natif), Anglais (courant, TOEIC 890)"])
    save(doc, "01_karim_bensalah_excellent.docx")


# ──────────────────────────────────────────────────────────────
# CV 2 — Nadia Lefebvre — BON MATCH (manque Terraform, AWS)
# ──────────────────────────────────────────────────────────────
def cv2():
    doc = Document()
    heading(doc, "Nadia LEFEBVRE")
    doc.add_paragraph("Ingénieure DevOps | Lyon → Paris | nadia.lefebvre@email.fr | +33 6 23 45 67 89")

    section(doc, "Résumé", [
        "Ingénieure DevOps avec 2 ans d'expérience, spécialisée dans la containerisation et l'automatisation CI/CD. "
        "Bonne maîtrise de Docker, Kubernetes et GitHub Actions. Travaille principalement sur Azure "
        "et souhaite évoluer vers AWS. Aucune expérience avec Terraform ou Ansible."
    ])

    section(doc, "Compétences techniques", [
        "Conteneurisation : Docker, Kubernetes, Docker Compose",
        "Cloud : Microsoft Azure (AKS, Blob Storage, Azure DevOps) — pas d'expérience AWS",
        "CI/CD : GitHub Actions, Azure Pipelines",
        "Scripting : Python (scripts d'automatisation), Bash",
        "Monitoring : Datadog, Azure Monitor",
        "OS : Linux (Ubuntu), administration de base",
        "Versioning : Git, GitHub",
    ])

    section(doc, "Expérience professionnelle", [
        "Ingénieure DevOps — InnovateSoft Lyon (Mars 2022 – Présent)",
        "- Mise en place de pipelines CI/CD GitHub Actions pour 5 applications",
        "- Déploiement et gestion de clusters Kubernetes sur Azure AKS",
        "- Containerisation de 3 applications Python/Node.js avec Docker",
        "- Scripts Python d'automatisation de déploiement",
        "",
        "Stagiaire DevOps — WebAgency (Fév 2022 – Mars 2022)",
        "- Initiation Docker et Linux",
        "- Support aux déploiements manuels",
    ])

    section(doc, "Formation", [
        "Master Informatique spécialité Cloud Computing — Université de Lyon (2022)",
        "Licence Informatique — IUT Lyon 1 (2020)",
    ])

    section(doc, "Langues", ["Français (natif), Anglais (intermédiaire, B2)"])
    save(doc, "02_nadia_lefebvre_bon.docx")


# ──────────────────────────────────────────────────────────────
# CV 3 — Sofiane Marouani — MOYEN (junior, compétences partielles)
# ──────────────────────────────────────────────────────────────
def cv3():
    doc = Document()
    heading(doc, "Sofiane MAROUANI")
    doc.add_paragraph("Technicien Informatique / Junior DevOps | Bordeaux | sofiane.m@email.fr")

    section(doc, "Résumé", [
        "Technicien informatique avec 18 mois d'expérience, en reconversion vers le DevOps. "
        "Connaissance de Docker acquise en autodidacte, notions de Linux et Git. "
        "Peu d'expérience CI/CD, pas de Kubernetes ni cloud professionnel."
    ])

    section(doc, "Compétences", [
        "Docker : déploiement de conteneurs simples, docker-compose basique",
        "Linux : Ubuntu, commandes de base, gestion de fichiers",
        "Git : commits, branches, GitHub",
        "Python : scripts simples, niveau intermédiaire",
        "Réseaux : notions TCP/IP, DNS, HTTP",
        "Windows Server : administration basique",
    ])

    section(doc, "Expérience", [
        "Technicien Support IT — PME Bordeaux (Jan 2023 – Présent)",
        "- Maintenance parc informatique (50 postes)",
        "- Installation logiciels, gestion utilisateurs Active Directory",
        "- Premiers déploiements Docker en environnement de développement",
        "",
        "Alternant Informatique — Collectivité locale (2021 – 2022)",
        "- Support utilisateurs niveau 1 et 2",
        "- Gestion tickets helpdesk",
    ])

    section(doc, "Formation", [
        "Licence Professionnelle Systèmes et Réseaux — IUT Bordeaux (2022)",
        "BTS SIO (Services Informatiques aux Organisations) (2020)",
    ])

    section(doc, "Langues", ["Français (natif), Anglais (scolaire, A2)"])
    save(doc, "03_sofiane_marouani_moyen.docx")


# ──────────────────────────────────────────────────────────────
# CV 4 — Emma Renard — FAIBLE (profil backend Java, aucun DevOps)
# ──────────────────────────────────────────────────────────────
def cv4():
    doc = Document()
    heading(doc, "Emma RENARD")
    doc.add_paragraph("Développeuse Backend Java | Paris | emma.renard@email.fr | +33 6 34 56 78 90")

    section(doc, "Résumé", [
        "Développeuse backend Java/Spring Boot avec 2 ans d'expérience dans le développement "
        "d'APIs REST et d'applications web. Aucune expérience DevOps, cloud ou infrastructure. "
        "Cherche à évoluer vers du développement full-stack."
    ])

    section(doc, "Compétences", [
        "Langages : Java 17, Spring Boot, Spring MVC, Hibernate, JPA",
        "Bases de données : MySQL, PostgreSQL (usage applicatif uniquement)",
        "API : REST, Swagger/OpenAPI, Postman",
        "Tests : JUnit 5, Mockito",
        "Frontend : notions HTML/CSS, Angular (débutant)",
        "Versioning : Git, GitLab",
        "Méthodes : Scrum, Jira",
    ])

    section(doc, "Expérience", [
        "Développeuse Java — FinTech Solutions Paris (Sept 2022 – Présent)",
        "- Développement d'APIs REST pour une plateforme de paiement",
        "- Optimisation requêtes SQL et refactoring Spring Boot",
        "- Revues de code, documentation Swagger",
        "",
        "Stagiaire Développement — ESN Paris (Fév 2022 – Août 2022)",
        "- Développement modules Java/Spring",
        "- Correction de bugs, tests unitaires JUnit",
    ])

    section(doc, "Formation", [
        "Master Génie Logiciel — Université Paris Saclay (2022)",
        "Licence Informatique — Université Paris Sud (2020)",
    ])

    section(doc, "Langues", ["Français (natif), Anglais (B1)"])
    save(doc, "04_emma_renard_faible.docx")


# ──────────────────────────────────────────────────────────────
# CV 5 — Mehdi Oukili — INCOHÉRENT (se dit expert, texte dit débutant)
# ──────────────────────────────────────────────────────────────
def cv5():
    doc = Document()
    heading(doc, "Mehdi OUKILI")
    doc.add_paragraph("Ingénieur Systèmes & DevOps Expert | Marseille | mehdi.oukili@email.fr")

    section(doc, "Résumé", [
        "Expert Kubernetes et Terraform avec 3 ans d'expérience. Maîtrise avancée AWS et Ansible. "
        "Spécialiste CI/CD et automatisation DevOps. Fort niveau Docker et orchestration de conteneurs."
    ])

    section(doc, "Compétences déclarées", [
        "Kubernetes : expert (gestion de clusters de production)",
        "Terraform : expert (infrastructure as code complète)",
        "AWS : expert (EC2, S3, EKS, Lambda, IAM)",
        "Ansible : avancé",
        "Docker : avancé",
        "CI/CD : Jenkins, GitLab CI",
    ])

    section(doc, "Expérience professionnelle", [
        "Administrateur Systèmes — Entreprise Régionale Marseille (Juin 2021 – Présent)",
        "- Administration de serveurs Linux (Ubuntu, RedHat) pour 80 utilisateurs",
        "- Gestion sauvegardes et monitoring basique avec Nagios",
        "- Quelques expérimentations avec Docker en local (docker run, notions de base)",
        "- Découverte de Kubernetes via tutoriels en ligne, pas encore en production",
        "- Initiation à AWS : création de quelques instances EC2 pour des tests personnels",
        "- Notions de Terraform : suivi d'un cours Udemy (jamais utilisé en entreprise)",
        "",
        "Technicien Réseaux — ISP Marseille (Jan 2020 – Mai 2021)",
        "- Maintenance réseau LAN/WAN",
        "- Support niveau 2, configuration routeurs Cisco",
        "- Installation postes de travail Windows",
    ])

    section(doc, "Formation", [
        "BTS SIO Réseaux — Lycée Technique Marseille (2019)",
        "Pas de formation supérieure en informatique cloud ou DevOps",
    ])

    section(doc, "Langues", ["Français (natif), Anglais (A2 — très limité)"])
    save(doc, "05_mehdi_oukili_incoherent.docx")


if __name__ == "__main__":
    print("Génération des CV de test — Scénario DevOps/Cloud")
    print(f"Dossier de sortie : {os.path.abspath(OUT_DIR)}\n")
    cv1()
    cv2()
    cv3()
    cv4()
    cv5()
    print("\nTerminé. 5 CV générés.")
    print("\n" + "="*60)
    print("OFFRE À CRÉER DANS L'INTERFACE :")
    print("="*60)
    print("Titre      : Ingénieur DevOps / Cloud")
    print("Contrat    : CDI")
    print("Localisation : Paris")
    print("Postes     : 2")
    print("Compétences : Docker, Kubernetes, Terraform, AWS, CI/CD, Ansible, Linux, Python")
    print("Description :")
    print("""  Nous recherchons un ingénieur DevOps/Cloud expérimenté pour rejoindre notre équipe
  infrastructure. Vous serez responsable de la gestion et de l'évolution de nos
  environnements cloud AWS, de l'orchestration de nos conteneurs Kubernetes,
  et de l'automatisation de nos pipelines CI/CD. Vous maîtrisez Terraform pour
  l'infrastructure as code et Ansible pour l'automatisation. Expérience minimum
  3 ans en environnement DevOps/Cloud professionnel requise.
  Formation Bac+5 ingénierie informatique ou équivalent.""")
