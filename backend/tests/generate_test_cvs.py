"""
Générateur de CVs de test pour validation de qualité d'extraction
Crée des CVs dans différents formats avec contenu connu

Usage:
    python backend/tests/generate_test_cvs.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm


def create_test_docx():
    """Crée un CV de test au format DOCX"""
    print("Création du CV test DOCX...")
    
    doc = Document()
    
    # Titre
    title = doc.add_heading('CURRICULUM VITAE', 0)
    title.alignment = 1  # Centré
    
    # Informations personnelles
    doc.add_heading('Marie Martin', 1)
    doc.add_paragraph('Email: marie.martin@example.com')
    doc.add_paragraph('Téléphone: +33 6 12 34 56 78')
    doc.add_paragraph('LinkedIn: linkedin.com/in/marie-martin')
    
    # Profil
    doc.add_heading('Profil', 2)
    doc.add_paragraph(
        'Ingénieure DevOps passionnée avec 5 ans d\'expérience en automatisation, '
        'CI/CD et infrastructure cloud. Expertise en Python, Docker et Kubernetes.'
    )
    
    # Compétences
    doc.add_heading('Compétences', 2)
    skills = doc.add_paragraph()
    skills.add_run('• Langages: ').bold = True
    skills.add_run('Python, Bash, YAML, JavaScript\n')
    
    skills = doc.add_paragraph()
    skills.add_run('• DevOps: ').bold = True
    skills.add_run('Docker, Kubernetes, Jenkins, GitLab CI, Terraform\n')
    
    skills = doc.add_paragraph()
    skills.add_run('• Cloud: ').bold = True
    skills.add_run('AWS, Azure, GCP\n')
    
    skills = doc.add_paragraph()
    skills.add_run('• Bases de données: ').bold = True
    skills.add_run('PostgreSQL, MongoDB, Redis\n')
    
    # Expérience
    doc.add_heading('Expérience Professionnelle', 2)
    
    job1 = doc.add_heading('Ingénieure DevOps Senior - TechCorp', 3)
    doc.add_paragraph('2021 - Présent | Paris, France')
    doc.add_paragraph('• Mise en place pipeline CI/CD avec GitLab pour 20+ microservices')
    doc.add_paragraph('• Migration infrastructure vers Kubernetes (AWS EKS)')
    doc.add_paragraph('• Réduction temps déploiement de 4h à 15min')
    doc.add_paragraph('• Automatisation monitoring avec Prometheus et Grafana')
    
    job2 = doc.add_heading('Ingénieure DevOps - StartupXYZ', 3)
    doc.add_paragraph('2019 - 2021 | Lyon, France')
    doc.add_paragraph('• Infrastructure as Code avec Terraform')
    doc.add_paragraph('• Conteneurisation applications legacy avec Docker')
    doc.add_paragraph('• Mise en place logging centralisé (ELK Stack)')
    
    # Formation
    doc.add_heading('Formation', 2)
    doc.add_paragraph('Diplôme d\'Ingénieur en Informatique')
    doc.add_paragraph('École Centrale - 2019')
    
    # Langues
    doc.add_heading('Langues', 2)
    doc.add_paragraph('• Français: Langue maternelle')
    doc.add_paragraph('• Anglais: Courant (TOEIC 950)')
    doc.add_paragraph('• Allemand: Intermédiaire')
    
    # Sauvegarder
    output_path = Path('data/cvs_raw/test_word.docx')
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    
    print(f"✅ CV DOCX créé: {output_path.resolve()}")
    return output_path


def create_test_pdf():
    """Crée un CV de test au format PDF texte"""
    print("Création du CV test PDF...")
    
    output_path = Path('data/cvs_raw/test_pdf_text.pdf')
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Créer le PDF
    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4
    
    # Marges
    left_margin = 2.5 * cm
    top_margin = height - 2.5 * cm
    y = top_margin
    
    # Titre
    c.setFont("Helvetica-Bold", 20)
    c.drawCentredString(width / 2, y, "CURRICULUM VITAE")
    y -= 1.5 * cm
    
    # Nom
    c.setFont("Helvetica-Bold", 16)
    c.drawString(left_margin, y, "Jean Dupont")
    y -= 0.8 * cm
    
    # Informations de contact
    c.setFont("Helvetica", 11)
    c.drawString(left_margin, y, "Email: jean.dupont@example.com")
    y -= 0.6 * cm
    c.drawString(left_margin, y, "Téléphone: +33 6 98 76 54 32")
    y -= 0.6 * cm
    c.drawString(left_margin, y, "LinkedIn: linkedin.com/in/jean-dupont")
    y -= 1.2 * cm
    
    # Profil
    c.setFont("Helvetica-Bold", 14)
    c.drawString(left_margin, y, "Profil")
    y -= 0.7 * cm
    
    c.setFont("Helvetica", 11)
    profile_text = (
        "Développeur Full Stack avec 3 ans d'expérience en Python et React. "
        "Passionné par les architectures modernes et les bonnes pratiques de développement."
    )
    # Découper le texte pour qu'il tienne sur plusieurs lignes
    words = profile_text.split()
    line = ""
    for word in words:
        test_line = line + word + " "
        if c.stringWidth(test_line, "Helvetica", 11) < width - 2 * left_margin:
            line = test_line
        else:
            c.drawString(left_margin, y, line)
            y -= 0.5 * cm
            line = word + " "
    if line:
        c.drawString(left_margin, y, line)
    y -= 1 * cm
    
    # Compétences
    c.setFont("Helvetica-Bold", 14)
    c.drawString(left_margin, y, "Compétences")
    y -= 0.7 * cm
    
    c.setFont("Helvetica", 11)
    skills = [
        "• Langages: Python, JavaScript, TypeScript, SQL",
        "• Backend: FastAPI, Django, Node.js, PostgreSQL",
        "• Frontend: React, Vue.js, Tailwind CSS",
        "• DevOps: Docker, Git, CI/CD, AWS"
    ]
    for skill in skills:
        c.drawString(left_margin, y, skill)
        y -= 0.6 * cm
    y -= 0.6 * cm
    
    # Expérience
    c.setFont("Helvetica-Bold", 14)
    c.drawString(left_margin, y, "Expérience Professionnelle")
    y -= 0.7 * cm
    
    c.setFont("Helvetica-Bold", 12)
    c.drawString(left_margin, y, "Développeur Full Stack - WebAgency")
    y -= 0.6 * cm
    
    c.setFont("Helvetica-Oblique", 10)
    c.drawString(left_margin, y, "2023 - Présent | Paris, France")
    y -= 0.6 * cm
    
    c.setFont("Helvetica", 11)
    experiences = [
        "• Développement API REST avec FastAPI et SQLAlchemy",
        "• Interface admin React avec authentification JWT",
        "• Tests automatisés (Pytest, Cypress) - couverture 85%",
        "• Déploiement Docker sur AWS EC2"
    ]
    for exp in experiences:
        c.drawString(left_margin, y, exp)
        y -= 0.6 * cm
    y -= 0.4 * cm
    
    c.setFont("Helvetica-Bold", 12)
    c.drawString(left_margin, y, "Développeur Junior - StartupTech")
    y -= 0.6 * cm
    
    c.setFont("Helvetica-Oblique", 10)
    c.drawString(left_margin, y, "2021 - 2023 | Lyon, France")
    y -= 0.6 * cm
    
    c.setFont("Helvetica", 11)
    experiences2 = [
        "• Maintenance et évolution application Django",
        "• Intégration API tierces (Stripe, SendGrid)",
        "• Optimisation performances base de données"
    ]
    for exp in experiences2:
        c.drawString(left_margin, y, exp)
        y -= 0.6 * cm
    y -= 0.6 * cm
    
    # Formation
    c.setFont("Helvetica-Bold", 14)
    c.drawString(left_margin, y, "Formation")
    y -= 0.7 * cm
    
    c.setFont("Helvetica", 11)
    c.drawString(left_margin, y, "Master Informatique - Spécialité Génie Logiciel")
    y -= 0.6 * cm
    c.drawString(left_margin, y, "Université de Lyon - 2021")
    
    # Sauvegarder
    c.save()
    
    print(f"✅ CV PDF créé: {output_path.resolve()}")
    return output_path


def create_reference_texts():
    """Crée un fichier avec les textes de référence pour la validation"""
    print("Création du fichier de référence...")
    
    references = {
        "test_word.docx": """
CURRICULUM VITAE
Marie Martin
Email: marie.martin@example.com
Téléphone: +33 6 12 34 56 78
LinkedIn: linkedin.com/in/marie-martin

Profil
Ingénieure DevOps passionnée avec 5 ans d'expérience en automatisation,
CI/CD et infrastructure cloud. Expertise en Python, Docker et Kubernetes.

Compétences
Langages: Python, Bash, YAML, JavaScript
DevOps: Docker, Kubernetes, Jenkins, GitLab CI, Terraform
Cloud: AWS, Azure, GCP
Bases de données: PostgreSQL, MongoDB, Redis

Expérience Professionnelle
Ingénieure DevOps Senior - TechCorp
2021 - Présent | Paris, France
Mise en place pipeline CI/CD avec GitLab pour 20+ microservices
Migration infrastructure vers Kubernetes (AWS EKS)
Réduction temps déploiement de 4h à 15min
Automatisation monitoring avec Prometheus et Grafana

Formation
Diplôme d'Ingénieur en Informatique
École Centrale - 2019

Langues
Français: Langue maternelle
Anglais: Courant (TOEIC 950)
Allemand: Intermédiaire
        """,
        
        "test_pdf_text.pdf": """
CURRICULUM VITAE
Jean Dupont
Email: jean.dupont@example.com
Téléphone: +33 6 98 76 54 32
LinkedIn: linkedin.com/in/jean-dupont

Profil
Développeur Full Stack avec 3 ans d'expérience en Python et React.
Passionné par les architectures modernes et les bonnes pratiques de développement.

Compétences
Langages: Python, JavaScript, TypeScript, SQL
Backend: FastAPI, Django, Node.js, PostgreSQL
Frontend: React, Vue.js, Tailwind CSS
DevOps: Docker, Git, CI/CD, AWS

Expérience Professionnelle
Développeur Full Stack - WebAgency
2023 - Présent | Paris, France
Développement API REST avec FastAPI et SQLAlchemy
Interface admin React avec authentification JWT
Tests automatisés (Pytest, Cypress) - couverture 85%
Déploiement Docker sur AWS EC2

Formation
Master Informatique - Spécialité Génie Logiciel
Université de Lyon - 2021
        """
    }
    
    output_path = Path('data/cvs_raw/reference_texts.txt')
    with open(output_path, 'w', encoding='utf-8') as f:
        for filename, text in references.items():
            f.write(f"{'='*70}\n")
            f.write(f"RÉFÉRENCE: {filename}\n")
            f.write(f"{'='*70}\n")
            f.write(text.strip())
            f.write(f"\n\n")
    
    print(f"✅ Références créées: {output_path.resolve()}")
    return output_path


if __name__ == "__main__":
    print("""
╔═══════════════════════════════════════════════════════════════════╗
║         GÉNÉRATION CVs DE TEST - TalentMatch                      ║
║                     Sprint 1 - Validation Qualité                 ║
╚═══════════════════════════════════════════════════════════════════╝
    """)
    
    try:
        # Créer les CVs de test
        docx_path = create_test_docx()
        pdf_path = create_test_pdf()
        ref_path = create_reference_texts()
        
        print(f"\n{'='*70}")
        print("✅ TOUS LES FICHIERS DE TEST ONT ÉTÉ CRÉÉS")
        print(f"{'='*70}")
        print(f"\nFichiers créés:")
        print(f"  1. {docx_path.name} - CV DOCX (Marie Martin - DevOps)")
        print(f"  2. {pdf_path.name} - CV PDF texte (Jean Dupont - Full Stack)")
        print(f"  3. {ref_path.name} - Textes de référence pour validation")
        
        print(f"\n💡 Prochaine étape:")
        print(f"   Lancer le test de qualité:")
        print(f"   python backend/tests/test_extraction_quality.py")
        print(f"   OU")
        print(f"   pytest backend/tests/test_extraction_quality.py -v -s")
        
    except Exception as e:
        print(f"\n❌ Erreur lors de la génération: {e}")
        import traceback
        traceback.print_exc()
