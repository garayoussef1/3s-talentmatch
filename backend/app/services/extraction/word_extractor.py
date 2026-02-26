"""
Module d'extraction de texte depuis des fichiers Word (.docx)
Partie du projet 3S TalentMatch
"""

from docx import Document
from pathlib import Path
from typing import Dict
import logging

logger = logging.getLogger(__name__)


class WordExtractor:
    """Extracteur de texte pour fichiers Word (.docx)"""

    def __init__(self, min_text_length: int = 100):
        self.min_text_length = min_text_length
        logger.info("WordExtractor initialisé")

    def extract(self, file_path: str) -> Dict:
        """
        Extrait le texte d'un fichier Word (.docx)

        Returns:
            dict avec success, text, pages, method, needs_ocr, error
        """
        logger.info(f"Début extraction Word: {file_path}")

        result = {
            'success': False,
            'text': '',
            'pages': 0,
            'method': 'python-docx',
            'needs_ocr': False,
            'error': None
        }

        try:
            path = Path(file_path)

            # Vérifier existence
            if not path.exists():
                result['error'] = f"Fichier non trouvé: {file_path}"
                logger.error(result['error'])
                return result

            # Vérifier extension
            if path.suffix.lower() not in ['.docx']:
                result['error'] = f"Extension invalide: {path.suffix} (attendu: .docx)"
                logger.error(result['error'])
                return result

            # Lire le document
            doc = Document(str(path))

            # Extraire les paragraphes
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Extraire le texte des tableaux
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text and text not in paragraphs:
                            paragraphs.append(text)

            full_text = '\n'.join(paragraphs)

            # Word n'a pas de "pages" directement accessibles
            result['pages'] = 1
            result['text'] = full_text

            if len(full_text) < self.min_text_length:
                result['error'] = 'Document vide ou texte insuffisant'
                logger.warning(result['error'])
                return result

            result['success'] = True
            logger.info(f"Extraction Word réussie: {len(full_text)} caractères")
            return result

        except Exception as e:
            result['error'] = f"Erreur extraction Word: {str(e)}"
            logger.exception(result['error'])
            return result
